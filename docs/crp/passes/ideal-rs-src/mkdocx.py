# -*- coding: utf-8 -*-
"""Build the house-style .docx from rs_word.html, preserving figure text wrapping."""
import re, html as H
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TEAL = RGBColor(0x14,0x45,0x3F); RED = RGBColor(0x9B,0x1C,0x1C); BLACK = RGBColor(0,0,0)
SRC = open('rs_word.html', encoding='utf-8').read()
SRC = SRC.split('<body>',1)[1].rsplit('</body>',1)[0]

doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(8.5), Inches(11)
for m in ('top_margin','bottom_margin','left_margin','right_margin'):
    setattr(sec, m, Inches(0.5))
st = doc.styles['Normal']
st.font.name = 'Arial'; st.font.size = Pt(11); st.font.color.rgb = BLACK
st.element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
pf = st.paragraph_format
pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
pf.space_after = Pt(3); pf.space_before = Pt(0)
pf.line_spacing_rule = WD_LINE_SPACING.SINGLE

def border(p_or_tc, edge='bottom', sz=8, color='000000'):
    pPr = p_or_tc.paragraph_format.element.get_or_add_pPr()
    bdr = OxmlElement('w:pBdr'); e = OxmlElement('w:'+edge)
    e.set(qn('w:val'),'single'); e.set(qn('w:sz'),str(sz))
    e.set(qn('w:space'),'2'); e.set(qn('w:color'),color)
    bdr.append(e); pPr.append(bdr)

def shade(cell, hexc):
    tcPr = cell._tc.get_or_add_tcPr(); sh = OxmlElement('w:shd')
    sh.set(qn('w:val'),'clear'); sh.set(qn('w:fill'),hexc); tcPr.append(sh)

def cell_border(cell, edge, sz, color):
    tcPr = cell._tc.get_or_add_tcPr()
    b = tcPr.find(qn('w:tcBorders'))
    if b is None: b = OxmlElement('w:tcBorders'); tcPr.append(b)
    e = OxmlElement('w:'+edge); e.set(qn('w:val'),'single')
    e.set(qn('w:sz'),str(sz)); e.set(qn('w:space'),'0'); e.set(qn('w:color'),color)
    b.append(e)

def float_table(tbl, width_in=None):
    """Right-float the table with square text wrap, the Word-native equivalent
    of the PDF's floated figure."""
    tblPr = tbl._tbl.tblPr
    p = OxmlElement('w:tblpPr')
    for k,v in [('w:leftFromText','180'),('w:rightFromText','0'),
                ('w:topFromText','60'),('w:bottomFromText','60'),
                ('w:vertAnchor','text'),('w:horzAnchor','margin'),
                ('w:tblpY','1')]:
        p.set(qn(k), v)
    # Position the frame explicitly. tblpXSpec="right" lands 0.13in outside the
    # margin in Writer, so compute the left edge from the 7.5in text column.
    if width_in is not None:
        # Writer adds ~0.13in to the computed x for a text-anchored frame.
        # Pull it back so the frame is inside the right margin in Writer, and
        # 0.13in short of it in Word. Erring inward is the safe direction.
        p.set(qn('w:tblpX'), str(int(round((7.5 - width_in - 0.13) * 1440))))
    else:
        p.set(qn('w:tblpXSpec'), 'right')
    tblPr.insert(0, p)
    ov = OxmlElement('w:tblOverlap'); ov.set(qn('w:val'),'never'); tblPr.append(ov)

def fixed_width(tbl, inches):
    """Pin the table width so Word does not autofit it back to the full column."""
    tblPr = tbl._tbl.tblPr
    for tag in ('w:tblW','w:tblLayout'):
        old = tblPr.find(qn(tag))
        if old is not None: tblPr.remove(old)
    w = OxmlElement('w:tblW'); w.set(qn('w:w'), str(int(inches*1440))); w.set(qn('w:type'),'dxa')
    tblPr.append(w)
    lay = OxmlElement('w:tblLayout'); lay.set(qn('w:type'),'fixed'); tblPr.append(lay)

def zero_cell_margins(tbl):
    """Word pads cells 0.08in each side by default, which would push a
    full-width image past the frame. Zero it for figure frames."""
    tblPr = tbl._tbl.tblPr
    m = OxmlElement('w:tblCellMar')
    for e in ('top','left','bottom','right'):
        x = OxmlElement('w:'+e); x.set(qn('w:w'),'0'); x.set(qn('w:type'),'dxa'); m.append(x)
    tblPr.append(m)

def no_borders(tbl):
    tblPr = tbl._tbl.tblPr
    b = OxmlElement('w:tblBorders')
    for e in ('top','left','bottom','right','insideH','insideV'):
        x = OxmlElement('w:'+e); x.set(qn('w:val'),'none'); x.set(qn('w:sz'),'0'); b.append(x)
    tblPr.append(b)

TOKEN = re.compile(r'(<b>|</b>|<i>|</i>|<sup>|</sup>|<br\s*/?>|<span class="eg">|</span>)', re.I)
def add_runs(par, frag, size=None, color=None, base_bold=False, base_ital=False):
    bold, ital, sup, eg = base_bold, base_ital, False, False
    for tok in TOKEN.split(frag):
        if not tok: continue
        t = tok.lower()
        if t == '<b>': bold = True
        elif t == '</b>': bold = base_bold
        elif t == '</span>': eg = False; ital = base_ital; bold = base_bold
        elif t == '<span class="eg">':
            # the matrix "for example" line: its own line, smaller and italic
            par.add_run().add_break(); eg = True; ital = True; bold = False
        elif t == '<i>': ital = True
        elif t == '</i>': ital = base_ital
        elif t == '<sup>': sup = True
        elif t == '</sup>': sup = False
        elif t.startswith('<br'): par.add_run().add_break()
        else:
            txt = H.unescape(re.sub(r'<[^>]+>','',tok))
            txt = re.sub(r'\s+',' ',txt)
            if not txt: continue
            r = par.add_run(txt)
            r.bold, r.italic = bold, ital
            if sup: r.font.superscript = True
            if eg: r.font.size = Pt(7.5)
            elif size: r.font.size = Pt(size)
            if color is not None: r.font.color.rgb = color
    return par

def strip_outer(s, tag):
    return re.sub(rf'^\s*<{tag}[^>]*>|</{tag}>\s*$', '', s.strip(), flags=re.I|re.S)

# ---- split the source into ordered blocks -------------------------------
BLOCK = re.compile(
  r'(<h1[^>]*>.*?</h1>|<table[^>]*>.*?</table>|<div class="fig">.*?</div>|'
  r'<div class="refs">.*?</div>|<p[^>]*>.*?</p>)', re.S)

def parse_table(thtml):
    rows = re.findall(r'<tr[^>]*>(.*?)</tr>', thtml, re.S)
    out = []
    for r in rows:
        cls = 'tot' if 'class="tot"' in thtml[:thtml.find(r)][-60:] else ''
        cells = re.findall(r'<(th|td)([^>]*)>(.*?)</\1>', r, re.S)
        out.append(cells)
    hdr = 1 if '<thead>' in thtml else 0
    return out, hdr

def widths_from(thtml):
    return [float(x) for x in re.findall(r'width:([\d.]+)%', thtml)] or None

def build_table(thtml, matrix=False):
    rows_raw = re.findall(r'<tr[^>]*>(.*?)</tr>', thtml, re.S)
    row_classes = re.findall(r'<tr([^>]*)>', thtml)
    grid = []
    for r in rows_raw:
        grid.append(re.findall(r'<(th|td)([^>]*)>(.*?)</\1>', r, re.S))
    ncol = max(len(r) for r in grid)
    tbl = doc.add_table(rows=len(grid), cols=ncol)
    tbl.autofit = False
    no_borders(tbl)
    fixed_width(tbl, 7.5)
    pct = widths_from(thtml)
    if pct:
        for ci,p in enumerate(pct[:ncol]):
            for row in tbl.rows:
                row.cells[ci].width = Inches(7.5*p/100.0)
    for ri,row in enumerate(grid):
        is_head = row and row[0][0].lower()=='th'
        rcls = row_classes[ri] if ri < len(row_classes) else ''
        for ci,(tag,attrs,content) in enumerate(row):
            cell = tbl.cell(ri,ci)
            cell.text = ''
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(1.5)
            p.paragraph_format.space_before = Pt(1.5)
            right = 'class="n"' in attrs or ('n"' in attrs and 'class="n' in attrs)
            centre = matrix and 'rowlab' not in attrs
            p.alignment = (WD_ALIGN_PARAGRAPH.CENTER if centre else
                           WD_ALIGN_PARAGRAPH.RIGHT if right else WD_ALIGN_PARAGRAPH.LEFT)
            own = 'own' in attrs
            col = None
            if is_head or 'rowlab' in attrs: col = TEAL
            if own: col = RGBColor(0xFF,0xFF,0xFF)
            add_runs(p, content, size=9, color=col,
                     base_bold=is_head or 'rowlab' in attrs or own)
            if own: shade(cell, '14453F')
            if 'class="tot"' in rcls: shade(cell, 'EEF3F0')
            if 'class="rem"' in rcls: shade(cell, 'FBEEEC')
            if is_head:
                cell_border(cell,'bottom',8,'14453F')
            elif ri == len(grid)-1:
                cell_border(cell,'bottom',8,'14453F')
            else:
                cell_border(cell,'bottom',2,'B9C4BD')
    return tbl

blocks = [b for b in BLOCK.split(SRC) if b.strip()]
for blk in blocks:
    b = blk.strip()
    if not b.startswith('<'): continue

    if b.startswith('<h1'):
        txt = re.sub(r'<[^>]+>','',b).strip().upper()
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(3)
        r = p.add_run(txt); r.bold = True
        border(p,'bottom',10,'000000')

    elif b.startswith('<table') and 'figfloat' in b:
        img = re.search(r'src="([^"]+)"[^>]*style="width:([\d.]+)in"', b)
        cap = re.search(r'<p class="caption">(.*?)</p>', b, re.S)
        w = float(img.group(2))
        t = doc.add_table(rows=1, cols=1); t.autofit = False
        t.columns[0].width = Inches(w); t.rows[0].cells[0].width = Inches(w)
        no_borders(t); zero_cell_margins(t); fixed_width(t, w); float_table(t, w)
        c = t.cell(0,0); c.text = ''
        pi = c.paragraphs[0]; pi.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pi.paragraph_format.space_after = Pt(2)
        pi.add_run().add_picture(img.group(1), width=Inches(w))
        pc = c.add_paragraph(); pc.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pc.paragraph_format.space_after = Pt(0)
        add_runs(pc, cap.group(1), size=9)

    elif b.startswith('<table'):
        build_table(b, matrix='class="matrix"' in b)

    elif b.startswith('<div class="fig">'):
        img = re.search(r'src="([^"]+)"[^>]*style="width:([\d.]+)in"', b)
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(2)
        p.add_run().add_picture(img.group(1), width=Inches(float(img.group(2))))
        p.paragraph_format.keep_with_next = True

    elif b.startswith('<div class="refs">'):
        doc.add_page_break()
        inner = b
        h = re.search(r'<h1[^>]*>(.*?)</h1>', inner, re.S)
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(4)
        p.add_run(re.sub(r'<[^>]+>','',h.group(1)).strip().upper()).bold = True
        border(p,'bottom',10,'000000')
        for m in re.finditer(r'<p>(.*?)</p>', inner, re.S):
            rp = doc.add_paragraph(); rp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            rp.paragraph_format.space_after = Pt(3)
            rp.paragraph_format.left_indent = Inches(0.2)
            rp.paragraph_format.first_line_indent = Inches(-0.2)
            add_runs(rp, m.group(1), size=9)

    elif b.startswith('<p'):
        cls = (re.search(r'class="([^"]*)"', b) or [None,''])[1]
        inner = strip_outer(b, 'p')
        if 'clearfix' in cls: continue
        p = doc.add_paragraph()
        if 'caption' in cls:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(5)
            add_runs(p, inner, size=9)
        elif 'tnote' in cls:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(5)
            add_runs(p, inner, size=9, base_ital=True)
        elif 'metrics-head' in cls:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(5); p.paragraph_format.space_after = Pt(2)
            r = add_runs(p, inner, base_ital=True).runs
            for rr in r: rr.underline = True
            p.paragraph_format.keep_with_next = True
        elif 'aimhead' in cls:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(9); p.paragraph_format.space_after = Pt(3)
            add_runs(p, inner, color=TEAL, base_bold=True)
            border(p,'bottom',6,'14453F')
            p.paragraph_format.keep_with_next = True
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            if 'sec' in cls.split():
                p.paragraph_format.space_before = Pt(7)
            add_runs(p, inner)

doc.save('Olera_CRP_ResearchStrategy.docx')
print('docx written')
