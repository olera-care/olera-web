# -*- coding: utf-8 -*-
"""Build the house-style .docx for the pre-CRP plan, iteration 4.

Differs from its siblings in three places: the week blocks are a
nested div that has to be flattened before the block split, the
span classes that carry the risk tags and the checkpoint markers
need run-level size and colour, and a quiet week greys entirely.
"""
import re, html as H
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TEAL = RGBColor(0x14, 0x45, 0x3F); BLACK = RGBColor(0, 0, 0)
GREY = RGBColor(0x5F, 0x6B, 0x64)
SRC = open('plan_word.html', encoding='utf-8').read()


def _match(html, start):
    """Index just past the </div> that closes the <div ...> beginning at start."""
    i, depth = start, 0
    while i < len(html):
        if html.startswith('<div', i):
            depth += 1; i += 4
        elif html.startswith('</div>', i):
            depth -= 1; i += 6
            if depth == 0:
                return i
        else:
            i += 1
    raise AssertionError('unclosed div')


def flatten(html):
    """Nested divs defeat a regex block split, so normalise them first.

    A figblk is only a page-break wrapper and can be dropped: the Word export
    keeps a figure with its caption through keep_with_next. A figwrap becomes a
    single flat marker carrying the float width, the image and the caption."""
    out = []
    i = 0
    while i < len(html):
        if html.startswith('<div class="refcols">', i):
            end = _match(html, i)
            out.append(flatten(html[i + len('<div class="refcols">'):end - 6]))
            i = end
        elif html.startswith('<div class="wk', i):
            j = html.index('>', i)
            quiet = 'quiet' in html[i:j]
            end = _match(html, i)
            inner = flatten(html[j + 1:end - 6])
            if quiet:
                inner = re.sub(r'class="(wkh|wkj|wkt|wko)"', r'class="\1 q"', inner)
            out.append(inner)
            i = end
        elif html.startswith('<div class="figblk">', i):
            end = _match(html, i)
            out.append(flatten(html[i + len('<div class="figblk">'):end - 6]))
            i = end
        elif html.startswith('<div class="figwrap"', i):
            end = _match(html, i)
            blk = html[i:end]
            w = re.search(r'width:([\d.]+)in', blk).group(1)
            img = re.search(r'<img[^>]*>', blk).group(0)
            cap = re.search(r'<p class="caption">(.*?)</p>', blk, re.S)
            out.append(f'<div class="figfloat" data-w="{w}">{img}'
                       f'<p class="caption">{cap.group(1) if cap else ""}</p></div>')
            i = end
        else:
            out.append(html[i]); i += 1
    return ''.join(out)


SRC = flatten(SRC)

doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(8.5), Inches(11)
for m in ('top_margin', 'bottom_margin', 'left_margin', 'right_margin'):
    setattr(sec, m, Inches(0.5))
st = doc.styles['Normal']
st.font.name = 'Arial'; st.font.size = Pt(11); st.font.color.rgb = BLACK
st.element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
pf = st.paragraph_format
pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
pf.space_after = Pt(3); pf.space_before = Pt(0)
pf.line_spacing_rule = WD_LINE_SPACING.SINGLE

def exact(p, pts):
    """Pin a paragraph's line box to the PDF's computed line height.

    Word's SINGLE rule is a font-metric height, roughly 1.15 em, while the PDF
    uses an explicit line-height. Left alone the two drift about half a point a
    line, which is enough to move a page break by the end of a section."""
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(pts)

def border(p, edge='bottom', sz=10, color='000000'):
    pPr = p.paragraph_format.element.get_or_add_pPr()
    bdr = OxmlElement('w:pBdr'); e = OxmlElement('w:' + edge)
    e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), str(sz))
    e.set(qn('w:space'), '2'); e.set(qn('w:color'), color)
    bdr.append(e); pPr.append(bdr)

def cant_split(row):
    """The stylesheet's `table.dat tr { break-inside: avoid }`. Without it Word
    splits a row across a page break and packs more onto the page than the print
    renderer does, which is what puts the two exports a row out of step."""
    trPr = row._tr.get_or_add_trPr()
    e = OxmlElement('w:cantSplit'); e.set(qn('w:val'), 'true'); trPr.append(e)

def no_borders(tbl):
    b = OxmlElement('w:tblBorders')
    for e in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        x = OxmlElement('w:' + e); x.set(qn('w:val'), 'none'); x.set(qn('w:sz'), '0'); b.append(x)
    tbl._tbl.tblPr.append(b)

def zero_cell_margins(tbl):
    m = OxmlElement('w:tblCellMar')
    for e in ('top', 'left', 'bottom', 'right'):
        x = OxmlElement('w:' + e); x.set(qn('w:w'), '0'); x.set(qn('w:type'), 'dxa'); m.append(x)
    tbl._tbl.tblPr.append(m)

def cell_margins(tbl, left=0, right=120):
    """Match the stylesheet's `td { padding: ... 6pt ... 0 }`: nothing on the
    left, 6pt on the right. Word's default is 0.08in on both sides, which
    narrows every column by 0.16in, wraps headers mid-word, and makes rows
    taller here than in the print render."""
    m = OxmlElement('w:tblCellMar')
    for edge, w in (('top', 0), ('left', left), ('bottom', 0), ('right', right)):
        x = OxmlElement('w:' + edge)
        x.set(qn('w:w'), str(w)); x.set(qn('w:type'), 'dxa'); m.append(x)
    tbl._tbl.tblPr.append(m)

def repeat_header(row):
    """The stylesheet's `thead { display: table-header-group }`: a table that
    breaks across pages repeats its header. Word needs w:tblHeader on the row."""
    trPr = row._tr.get_or_add_trPr()
    e = OxmlElement('w:tblHeader'); e.set(qn('w:val'), 'true'); trPr.append(e)

def fixed_width(tbl, inches):
    tblPr = tbl._tbl.tblPr
    for tag in ('w:tblW', 'w:tblLayout'):
        old = tblPr.find(qn(tag))
        if old is not None: tblPr.remove(old)
    w = OxmlElement('w:tblW'); w.set(qn('w:w'), str(int(inches * 1440))); w.set(qn('w:type'), 'dxa')
    tblPr.append(w)
    lay = OxmlElement('w:tblLayout'); lay.set(qn('w:type'), 'fixed'); tblPr.append(lay)

def float_table(tbl, width_in):
    """Right-float with square text wrap: the Word-native equivalent of the
    PDF's floated figure. Writer adds ~0.13in to a text-anchored frame's x, so
    compute the left edge inward from the 7.5in text column rather than using
    tblpXSpec='right', which lands outside the margin."""
    tblPr = tbl._tbl.tblPr
    p = OxmlElement('w:tblpPr')
    for k, v in [('w:leftFromText', '180'), ('w:rightFromText', '0'),
                 ('w:topFromText', '60'), ('w:bottomFromText', '60'),
                 ('w:vertAnchor', 'text'), ('w:horzAnchor', 'margin'), ('w:tblpY', '1')]:
        p.set(qn(k), v)
    p.set(qn('w:tblpX'), str(int(round((7.5 - width_in - 0.13) * 1440))))
    tblPr.insert(0, p)
    ov = OxmlElement('w:tblOverlap'); ov.set(qn('w:val'), 'never'); tblPr.append(ov)

TOKEN = re.compile(
    r'(<b[^>]*>|</b>|<i>|</i>|<sup>|</sup>|<span class="[^"]*">|</span>)', re.I)
# size, colour, bold for each span class used in the memo
SPANS = {'rk': (8.5, GREY, False), 'd': (None, GREY, False),
         'tag': (8, TEAL, True), 'tag q': (8, GREY, True)}
def add_runs(par, frag, size=None, color=None, base_bold=False, bold_color=None,
             base_ital=False):
    bold, ital, sup = base_bold, base_ital, False
    span_stack = []
    for tok in TOKEN.split(frag):
        if not tok: continue
        t = tok.lower()
        if t.startswith('<span'):
            cls = re.search(r'class="([^"]*)"', tok).group(1)
            span_stack.append(SPANS.get(cls, (None, None, False)))
            continue
        elif t == '</span>':
            if span_stack: span_stack.pop()
            continue
        if t.startswith('<b'): bold = True
        elif t == '</b>': bold = base_bold
        elif t == '<i>': ital = True
        elif t == '</i>': ital = base_ital
        elif t == '<sup>': sup = True
        elif t == '</sup>': sup = False
        else:
            txt = H.unescape(re.sub(r'<[^>]+>', '', tok))
            txt = re.sub(r'\s+', ' ', txt)
            if not txt: continue
            r = par.add_run(txt)
            r.bold, r.italic = bold, ital
            if sup:
                # w:vertAlign makes Word shrink the run to about 58%, which would
                # put the reference markers at 4.6pt. Set the size explicitly and
                # raise the baseline instead, so 8pt is what actually prints.
                r.font.size = Pt(8)
                pos = OxmlElement('w:position'); pos.set(qn('w:val'), '7')
                r._element.get_or_add_rPr().append(pos)
            if size: r.font.size = Pt(size)
            if color is not None: r.font.color.rgb = color
            elif bold and bold_color is not None: r.font.color.rgb = bold_color
            if span_stack:
                ssz, scol, sbold = span_stack[-1]
                if ssz: r.font.size = Pt(ssz)
                if scol is not None: r.font.color.rgb = scol
                if sbold: r.bold = True
    return par

def add_figure(img, w, cap, floated=False):
    """Floated figures need a table, because a frame is the only Word construct
    that wraps text. Full-width figures must NOT use one: Writer applies its own
    cell padding regardless of w:tblCellMar, which pushes a 7.2in image past the
    cell and crops its right edge."""
    if not floated:
        pi = doc.add_paragraph()
        pi.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # LibreOffice adds its own font leading under an inline image and ignores
        # an EXACTLY line height on the line holding it, so the block comes out a
        # different height than the PDF's. How much depends on the image, so the
        # gap between image and caption is set by measurement: 6pt for this one.
        # space_before is left at 0 because LibreOffice collapses it against the
        # preceding paragraph's space_after.
        pi.paragraph_format.space_before = Pt(0); pi.paragraph_format.space_after = Pt(6)
        pi.paragraph_format.keep_with_next = True
        pi.add_run().add_picture(img, width=Inches(w))
        pc = doc.add_paragraph()
        pc.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pc.paragraph_format.space_after = Pt(1)
        exact(pc, 10.44)
        add_runs(pc, cap, size=9)
        return
    t = doc.add_table(rows=1, cols=1); t.autofit = False
    t.columns[0].width = Inches(w); t.rows[0].cells[0].width = Inches(w)
    no_borders(t); zero_cell_margins(t); fixed_width(t, w); float_table(t, w)
    c = t.cell(0, 0); c.text = ''
    pi = c.paragraphs[0]; pi.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pi.paragraph_format.space_after = Pt(2); pi.paragraph_format.space_before = Pt(4)
    pi.add_run().add_picture(img, width=Inches(w))
    pc = c.add_paragraph(); pc.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pc.paragraph_format.space_after = Pt(0)
    add_runs(pc, cap, size=9)

def shade(cell, hexc):
    tcPr = cell._tc.get_or_add_tcPr(); sh = OxmlElement('w:shd')
    sh.set(qn('w:val'), 'clear'); sh.set(qn('w:fill'), hexc); tcPr.append(sh)

def cell_border(cell, edge, sz, color):
    tcPr = cell._tc.get_or_add_tcPr()
    b = tcPr.find(qn('w:tcBorders'))
    if b is None:
        b = OxmlElement('w:tcBorders'); tcPr.append(b)
    e = OxmlElement('w:' + edge); e.set(qn('w:val'), 'single')
    e.set(qn('w:sz'), str(sz)); e.set(qn('w:space'), '0'); e.set(qn('w:color'), color)
    b.append(e)

def build_table(thtml, keep=False, plain=False):
    rows = [re.findall(r'<(th|td)([^>]*)>(.*?)</\1>', r, re.S)
            for r in re.findall(r'<tr[^>]*>(.*?)</tr>', thtml, re.S)]
    pct = [float(x) for x in re.findall(r'width:([\d.]+)%', thtml)]
    ncol = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=ncol)
    tbl.autofit = False
    no_borders(tbl); fixed_width(tbl, 7.5); cell_margins(tbl)
    # Word lays a fixed-layout table out from tblGrid, so setting only the cell
    # widths leaves the grid at add_table's equal columns and the proportions
    # come out wrong. Set the column (gridCol) width as well.
    for ci, p in enumerate(pct[:ncol]):
        w = Inches(7.5 * p / 100.0)
        tbl.columns[ci].width = w
        for row in tbl.rows:
            row.cells[ci].width = w
    for row in tbl.rows:
        cant_split(row)
    if rows and rows[0] and rows[0][0][0].lower() == 'th':
        repeat_header(tbl.rows[0])
    tr_attrs = re.findall(r'<tr([^>]*)>', thtml)
    for ri, row in enumerate(rows):
        is_head = bool(row) and row[0][0].lower() == 'th'
        gate = ri < len(tr_attrs) and ('class="gate"' in tr_attrs[ri]
                                          or 'class="here"' in tr_attrs[ri])
        for ci, (tag, attrs, content) in enumerate(row):
            cell = tbl.cell(ri, ci); cell.text = ''
            par = cell.paragraphs[0]
            par.alignment = WD_ALIGN_PARAGRAPH.LEFT
            par.paragraph_format.space_before = Pt(2.8)
            par.paragraph_format.space_after = Pt(2.8)
            exact(par, 10.44)
            # matches the PDF's `table.dat td b { color: #14453f }`: the bold
            # lead-in is teal, the prose after it stays black
            # A bulleted cell is a div per item so each one can hang-indent;
            # everything else splits on <br> as before. Word has no <br> that
            # keeps its formatting cleanly, so both become one paragraph per line.
            tis = re.findall(r'<div class="ti">(.*?)</div>', content, re.S)
            segments = tis or re.split(r'<br\s*/?>', content, flags=re.I)
            for si, seg in enumerate(segments):
                target = par if si == 0 else cell.add_paragraph()
                if si:
                    target.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    target.paragraph_format.space_before = Pt(0)
                    target.paragraph_format.space_after = Pt(2.8)
                    exact(target, 10.44)
                if tis:
                    target.paragraph_format.space_after = Pt(1.5)
                    target.paragraph_format.left_indent = Inches(0.097)
                    target.paragraph_format.first_line_indent = Inches(-0.097)
                    seg = seg.lstrip('\u2022 ')
                    r = target.add_run('\u2022 ')
                    r.font.size = Pt(9); r.font.color.rgb = TEAL
                bb = is_head or (ci == 0 and content.lstrip().startswith('<b>'))
                add_runs(target, seg, size=9, color=TEAL if is_head else None,
                         base_bold=bb, bold_color=None if is_head else TEAL)
            # Word has no table-level break-inside:avoid. keep_with_next on every
            # row but the last is the equivalent, and is what stops Table 2
            # stranding a row on the previous page.
            if keep and ri < len(rows) - 1:
                par.paragraph_format.keep_with_next = True
            # A header row alone at the foot of a page is a stranded header, and
            # it is also where the two exports fall a row out of step: the print
            # renderer moves the header down with its first row, Word does not
            # unless it is told to.
            if is_head:
                par.paragraph_format.keep_with_next = True
            if 'class="tot"' in thtml and ri == len(rows) - 1:
                shade(cell, 'EEF3F0')
            if gate:
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'EEF3F1')
                cell._tc.get_or_add_tcPr().append(shd)
            if plain:
                cell_border(cell, 'top', 6, 'B9C4BD')
            elif is_head or ri == len(rows) - 1:
                cell_border(cell, 'bottom', 8, '14453F')
            else:
                cell_border(cell, 'bottom', 2, 'B9C4BD')
    return tbl

BLOCK = re.compile(r'(<h1[^>]*>.*?</h1>|<h2[^>]*>.*?</h2>|<ul class="pts">.*?</ul>|<div class="figfloat"[^>]*>.*?</div>|'
                   r'<div class="fig">.*?</div>|<div class="chain">.*?</div>|'
                   r'<ol[^>]*>.*?</ol>|'
                   r'<table[^>]*>.*?</table>|<p[^>]*>.*?</p>)', re.S)

# A caption that follows a full-width <div class="fig"> belongs inside that
# figure's frame, so it can never be stranded from its image by a page break.
blocks = [b.strip() for b in BLOCK.split(SRC) if b.strip().startswith('<')]
i = 0
while i < len(blocks):
    b = blocks[i]

    if b.startswith('<h1'):
        txt = H.unescape(re.sub(r'<[^>]+>', '', b)).strip()
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(txt); r.bold = True; r.font.size = Pt(15); r.font.color.rgb = TEAL
        i += 1
        continue

    if b.startswith('<h2'):
        txt = H.unescape(re.sub(r'<[^>]+>', '', b)).strip().upper()
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if 'refhead' in b:
            p.paragraph_format.page_break_before = True
        p.paragraph_format.space_before = Pt(11); p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(txt); r.bold = True
        border(p, 'bottom', 12, '000000')

    elif b.startswith('<div class="figfloat"'):
        img = re.search(r'src="([^"]+)"', b)
        w = float(re.search(r'data-w="([\d.]+)"', b).group(1))
        cap = re.search(r'<p class="caption">(.*?)</p>', b, re.S)
        add_figure(img.group(1), w, cap.group(1), floated=True)

    elif b.startswith('<div class="chain">'):
        img = re.search(r'src="([^"]+)"[^>]*style="width:([\d.]+)in"', b)
        pi = doc.add_paragraph(); pi.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pi.paragraph_format.space_before = Pt(4); pi.paragraph_format.space_after = Pt(6)
        pi.add_run().add_picture(img.group(1), width=Inches(float(img.group(2))))

    elif b.startswith('<div class="fig">'):
        img = re.search(r'src="([^"]+)"[^>]*style="width:([\d.]+)in"', b)
        cap = ''
        if i + 1 < len(blocks) and 'class="caption"' in blocks[i + 1]:
            cap = re.search(r'<p class="caption">(.*?)</p>', blocks[i + 1], re.S).group(1)
            i += 1
        add_figure(img.group(1), float(img.group(2)), cap, floated=False)

    elif b.startswith('<table'):
        build_table(b, keep='dat keep' in b, plain='class="adv"' in b)
        if i + 1 < len(blocks) and 'class="caption"' in blocks[i + 1]:
            cap = re.search(r'<p class="caption">(.*?)</p>', blocks[i + 1], re.S).group(1)
            i += 1
            pc = doc.add_paragraph()
            pc.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pc.paragraph_format.space_before = Pt(2)
            pc.paragraph_format.space_after = Pt(6)
            add_runs(pc, cap, size=9)

    elif b.startswith('<ul class="pts">'):
        for li in re.findall(r'<li>(.*?)</li>', b, re.S):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.left_indent = Inches(0.26)
            p.paragraph_format.first_line_indent = Inches(-0.16)
            p.paragraph_format.space_after = Pt(3)
            exact(p, 13.2)
            r = p.add_run('\u2022  '); r.font.color.rgb = TEAL; r.bold = True
            add_runs(p, li)

    elif b.startswith('<ol'):
        # Explicit numerals with a hanging indent, rather than a Word list
        # style: pasted into another document, a real list renumbers against
        # whatever list precedes it. Plain runs cannot.
        wkt = 'class="wkt' in b
        qs = 'class="qs"' in b
        q = 'wkt q' in b
        items = re.findall(r'<li>(.*?)</li>', b, re.S)
        for k, li in enumerate(items, 1):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.left_indent = Inches(0.46 if wkt else 0.30)
            p.paragraph_format.first_line_indent = Inches(-0.20 if wkt else -0.24)
            p.paragraph_format.space_after = Pt(1 if wkt else 2)
            p.paragraph_format.space_before = Pt(0)
            exact(p, 12 if (wkt or qs) else 13.2)
            if wkt:
                # the whole week block stays together, milestone included
                p.paragraph_format.keep_with_next = True
            r = p.add_run('%d. ' % k)
            r.font.color.rgb = GREY if wkt else TEAL
            if wkt or qs: r.font.size = Pt(10)
            if not wkt: r.bold = True
            m = re.match(r'\s*<b class="rk">(.*?)</b>\s*(.*)$', li, re.S)
            if m:
                r = p.add_run(m.group(1) + ' '); r.bold = True; r.font.color.rgb = TEAL
                add_runs(p, m.group(2), size=10 if (wkt or qs) else None)
            else:
                add_runs(p, li, size=10 if (wkt or qs) else None,
                         color=GREY if q else None)

    elif b.startswith('<p class="clearfix"'):
        pass

    elif b.startswith('<p'):
        txt = re.sub(r'^\s*<p[^>]*>|</p>\s*$', '', b, flags=re.S)
        cls = re.search(r'<p class="([^"]*)"', b)
        cls = cls.group(1) if cls else ''
        p = doc.add_paragraph()
        q = ' q' in cls
        if 'wkh' in cls:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            # Word sets single-spaced 10pt tighter than the PDF's 1.2 line box,
            # so the block spacing is opened up to keep the two in step
            p.paragraph_format.space_before = Pt(8.5)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.keep_with_next = True
            exact(p, 12)
            add_runs(p, txt, size=10, color=GREY if q else None,
                     bold_color=GREY if q else TEAL)
        elif 'wkj' in cls:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.keep_with_next = True
            exact(p, 12)
            add_runs(p, txt, size=10, color=GREY if q else TEAL, base_ital=True)
        elif 'wko' in cls:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(4)
            exact(p, 10.8)
            add_runs(p, txt, size=9, color=GREY)
            if p.runs: p.runs[0].font.color.rgb = TEAL
        elif 'note' in cls:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(5)
            p.paragraph_format.space_after = Pt(3)
            exact(p, 11.4)
            add_runs(p, txt, size=9.5, color=GREY)
            if p.runs: p.runs[0].font.color.rgb = TEAL
        elif 'lede' in cls:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(9)
            exact(p, 12)
            add_runs(p, txt, size=10, color=RGBColor(0x5F, 0x6B, 0x64))
        elif 'aim' in cls:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.keep_with_next = True
            add_runs(p, txt, base_bold=True, color=TEAL)
        elif 'refs' in cls:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.left_indent = Inches(0.22)
            p.paragraph_format.first_line_indent = Inches(-0.22)
            p.paragraph_format.space_after = Pt(4)
            add_runs(p, txt)
        elif 'modnote' in cls:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(5)
            add_runs(p, txt, size=9, base_ital=True)
        elif 'caption' in cls:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(6)
            exact(p, 10.44)
            add_runs(p, txt, size=9)
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_before = Pt(5 if 'sec' in cls and 'first-sec' not in cls else 0)
            p.paragraph_format.space_after = Pt(3)
            exact(p, 13.2)
            # Word and the print renderer round a justified line differently, so a
            # paragraph can wrap to one more line here than there. With widow
            # control on, that one line drags two more onto the next page and the
            # two exports fall out of step; off, each page simply fills.
            p.paragraph_format.widow_control = False
            add_runs(p, txt)
    i += 1

doc.save('Olera_Pre_CRP_RD_Commercialization_and_Execution_Plan.docx')
print('saved docx')
